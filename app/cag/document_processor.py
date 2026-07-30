"""Document processing for the CAG knowledge cache.

Extract -> preserve structure -> clean -> segment into logical sections.

This is deliberately NOT vector chunking. Sections keep their heading context so
the knowledge cache stays human-readable and directly injectable into the model
context. Supported: PDF, DOCX, TXT, MD, XLSX, CSV.
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, List

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".xlsx", ".csv"}


@dataclass
class Section:
    """One logical, self-describing piece of a document."""

    heading: str
    text: str
    metadata: Dict[str, object] = field(default_factory=dict)

    def rendered(self) -> str:
        if self.heading:
            return f"## {self.heading}\n{self.text}"
        return self.text


# ---------------------------------------------------------------------------
# Cleaning
# ---------------------------------------------------------------------------
def clean_text(text: str) -> str:
    if not text:
        return ""
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    # Join words split across line breaks by PDF extraction ("well-\nbeing").
    text = re.sub(r"(\w)-\n(\w)", r"\1\2", text)
    # Drop page-number-only lines.
    text = re.sub(r"(?m)^\s*\d{1,3}\s*$", "", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


_HEADING_MAX_WORDS = 12


def _looks_like_heading(line: str) -> bool:
    """Heuristic heading detection for plain-text/PDF lines."""
    s = line.strip()
    if not s or len(s) > 90:
        return False
    words = s.split()
    if len(words) > _HEADING_MAX_WORDS:
        return False
    if s.endswith((".", ",", ";", "?", "!")):
        return False
    # ALL CAPS, Title Case, numbered heading, or markdown heading.
    if s.isupper() and len(words) >= 1:
        return True
    if re.match(r"^#{1,6}\s+\S", s):
        return True
    if re.match(r"^\d+(\.\d+)*[.)]?\s+[A-Z]", s):
        return True
    letters = [w for w in words if w[:1].isalpha()]
    if letters and sum(1 for w in letters if w[:1].isupper()) / len(letters) >= 0.75:
        return True
    return False


def _segment_by_headings(text: str, base_meta: Dict[str, object]) -> List[Section]:
    """Split cleaned text into sections using detected headings."""
    lines = text.split("\n")
    sections: List[Section] = []
    heading = ""
    buffer: List[str] = []

    def flush():
        body = "\n".join(buffer).strip()
        if body or heading:
            sections.append(Section(heading=heading, text=body, metadata=dict(base_meta)))
        buffer.clear()

    for line in lines:
        stripped = line.strip()
        if not stripped:
            buffer.append("")
            continue
        if _looks_like_heading(stripped):
            if buffer and any(b.strip() for b in buffer):
                flush()
            heading = re.sub(r"^#{1,6}\s*", "", stripped)
        else:
            buffer.append(stripped)
    flush()

    # Merge tiny fragments into the previous section to avoid noise.
    merged: List[Section] = []
    for sec in sections:
        if merged and len(sec.text) < 80 and not sec.heading:
            merged[-1].text = (merged[-1].text + "\n" + sec.text).strip()
        else:
            merged.append(sec)
    return [s for s in merged if s.text.strip() or s.heading.strip()]


# ---------------------------------------------------------------------------
# Format-specific extraction
# ---------------------------------------------------------------------------
def _from_pdf(path: Path, meta: Dict[str, object]) -> List[Section]:
    import fitz

    doc = fitz.open(path)
    try:
        pages = []
        for i in range(doc.page_count):
            pages.append((i + 1, doc[i].get_text("text")))
    finally:
        doc.close()

    sections: List[Section] = []
    for page_no, raw in pages:
        cleaned = clean_text(raw)
        if not cleaned:
            continue
        page_meta = dict(meta)
        page_meta["page"] = page_no
        sections.extend(_segment_by_headings(cleaned, page_meta))
    return sections


def _from_docx(path: Path, meta: Dict[str, object]) -> List[Section]:
    import docx

    document = docx.Document(str(path))
    sections: List[Section] = []
    heading = ""
    buffer: List[str] = []

    def flush():
        body = "\n".join(buffer).strip()
        if body or heading:
            sections.append(Section(heading=heading, text=body, metadata=dict(meta)))
        buffer.clear()

    for para in document.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = (para.style.name or "").lower() if para.style else ""
        if style.startswith("heading") or style == "title":
            if buffer:
                flush()
            heading = text
        else:
            buffer.append(text)
    flush()

    # Tables: keep row structure readable.
    for t_i, table in enumerate(document.tables, start=1):
        rows = []
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            sections.append(Section(heading=f"Table {t_i}", text="\n".join(rows),
                                    metadata=dict(meta)))
    return sections


def _from_xlsx(path: Path, meta: Dict[str, object]) -> List[Section]:
    from openpyxl import load_workbook

    wb = load_workbook(path, data_only=True, read_only=True)
    sections: List[Section] = []
    try:
        for sheet in wb.worksheets:
            rows = list(sheet.iter_rows(values_only=True))
            if not rows:
                continue
            headers = [str(h).strip() if h is not None else f"col_{i+1}"
                       for i, h in enumerate(rows[0])]
            lines = []
            for row in rows[1:]:
                pairs = []
                for i, cell in enumerate(row):
                    if cell is None:
                        continue
                    val = re.sub(r"\s+", " ", str(cell).strip())
                    if val:
                        pairs.append(f"{headers[i] if i < len(headers) else f'col_{i+1}'}: {val}")
                if pairs:
                    lines.append(" | ".join(pairs))
            if lines:
                s_meta = dict(meta)
                s_meta["sheet"] = sheet.title
                sections.append(Section(heading=sheet.title, text="\n".join(lines),
                                        metadata=s_meta))
    finally:
        wb.close()
    return sections


def _from_csv(path: Path, meta: Dict[str, object]) -> List[Section]:
    import pandas as pd

    df = pd.read_csv(path)
    lines = []
    for _, row in df.iterrows():
        pairs = [f"{c}: {row[c]}" for c in df.columns if not pd.isna(row[c])]
        if pairs:
            lines.append(" | ".join(pairs))
    if not lines:
        return []
    return [Section(heading=path.stem, text="\n".join(lines), metadata=dict(meta))]


def _from_text(path: Path, meta: Dict[str, object]) -> List[Section]:
    raw = path.read_text(encoding="utf-8", errors="replace")
    return _segment_by_headings(clean_text(raw), meta)


_EXTRACTORS = {
    ".pdf": _from_pdf,
    ".docx": _from_docx,
    ".xlsx": _from_xlsx,
    ".csv": _from_csv,
    ".txt": _from_text,
    ".md": _from_text,
}


def process_document(path: Path, knowledge_type: str = "general") -> List[Section]:
    """Extract a document into clean, structured sections."""
    ext = path.suffix.lower()
    extractor = _EXTRACTORS.get(ext)
    if extractor is None:
        return []
    meta: Dict[str, object] = {
        "document": path.name,
        "knowledge_type": knowledge_type,
        "format": ext.lstrip("."),
    }
    try:
        return extractor(path, meta)
    except Exception:
        # A single bad document must never break ingestion of the rest.
        return []
